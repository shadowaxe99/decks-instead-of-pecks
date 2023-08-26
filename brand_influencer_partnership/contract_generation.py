```python
from docx import Document
from docx.shared import Inches
from contract_sdk import ContractSDK

def generate_contract(influencer_id):
    # Fetch influencer data
    influencer_data = get_influencer_data(influencer_id)

    # Initialize contract SDK
    contract_sdk = ContractSDK()

    # Create a new Document
    doc = Document()

    # Add a title
    doc.add_heading('Contract Agreement', 0)

    # Add influencer details
    doc.add_paragraph('This contract is between InfluenceAI and ' + influencer_data['name'])

    # Add contract details
    doc.add_paragraph('Influencer agrees to promote the brand in exchange for ...')

    # Save the contract
    doc.save('contracts/' + influencer_id + '.docx')

    # Generate digital signature
    digital_signature = contract_sdk.generate_signature(influencer_data)

    return digital_signature

def get_influencer_data(influencer_id):
    # This function will fetch influencer data from the database
    # For the purpose of this code, we will return a dummy data
    return {
        'name': 'John Doe',
        'social_media_handle': '@johndoe',
        'engagement_rate': 0.8,
        'followers_count': 10000
    }
```